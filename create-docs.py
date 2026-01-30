#!/usr/bin/env python3
"""
Create One Development branded Word documents for Laguna Residence Investor Pitch
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from datetime import datetime

# Load deal data
with open('deal-data.json', 'r') as f:
    deal = json.load(f)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_aed(amount):
    """Format number as AED"""
    return f"AED {amount:,.0f}"

def create_external_proposal():
    """Create External Investor Proposal in One Development format"""
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('Investment Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Presented by One International Real Estate Development LLC')
    run.bold = True
    run.font.size = Pt(12)
    
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run('(All figures are indicative and subject to confirmation upon investor deal finalization)')
    run.italic = True
    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # 1. Overview
    doc.add_heading('1. Overview', level=1)
    
    overview = doc.add_paragraph()
    overview.add_run('Total Proposed Investment: ').bold = True
    overview.add_run(format_aed(deal['totals']['grossValue']) + '\n')
    overview.add_run('Structure: ').bold = True
    overview.add_run('Bulk Acquisition with Referral Commission\n')
    overview.add_run('Payment Terms: ').bold = True
    overview.add_run('100% Upfront Cash')
    
    # Overview table
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Component', 'Investment', 'Net Investment', 'Discount']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    row = table.rows[1].cells
    row[0].text = 'Laguna Residence Portfolio'
    row[1].text = format_aed(deal['totals']['grossValue'])
    row[2].text = format_aed(deal['totals']['investorPays'])
    row[3].text = '19% (10% Bulk + 9% Commission)'
    
    doc.add_paragraph()
    
    # 2. Project Overview
    doc.add_heading('2. Laguna Residence Investment', level=1)
    
    doc.add_heading('Project Overview', level=2)
    
    project_info = doc.add_paragraph()
    project_info.add_run('Project: ').bold = True
    project_info.add_run('Laguna Residence\n')
    project_info.add_run('Location: ').bold = True
    project_info.add_run('Dubai – UAE\n')
    project_info.add_run('Developer: ').bold = True
    project_info.add_run('One International Real Estate Development LLC\n')
    project_info.add_run('Total Units: ').bold = True
    project_info.add_run(str(deal['totals']['totalUnits']) + '\n')
    project_info.add_run('Unit Types: ').bold = True
    project_info.add_run('Studio – 3 Bedroom\n')
    project_info.add_run('Size Range: ').bold = True
    project_info.add_run('1,500 – 3,800 sqft\n')
    project_info.add_run('Handover: ').bold = True
    project_info.add_run('Q4 2027\n')
    project_info.add_run('Payment Plan: ').bold = True
    project_info.add_run('50/50')
    
    # Unit Mix Table
    doc.add_heading('Unit Portfolio', level=2)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Unit Type', 'Units', 'Avg. Price (AED)', 'Value (AED)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    # Data rows
    unit_data = [
        ('Studio', deal['unitMix']['studio']['count'], deal['unitMix']['studio']['avgPrice'], deal['unitMix']['studio']['subtotal']),
        ('1 Bedroom', deal['unitMix']['oneBedroom']['count'], deal['unitMix']['oneBedroom']['avgPrice'], deal['unitMix']['oneBedroom']['subtotal']),
        ('2 Bedroom', deal['unitMix']['twoBedroom']['count'], deal['unitMix']['twoBedroom']['avgPrice'], deal['unitMix']['twoBedroom']['subtotal']),
    ]
    
    for i, (unit_type, count, avg, total) in enumerate(unit_data, 1):
        row = table.rows[i].cells
        row[0].text = unit_type
        row[1].text = str(count)
        row[2].text = f"{avg:,.0f}"
        row[3].text = f"{total:,.0f}"
    
    # Total row
    total_row = table.rows[4].cells
    total_row[0].text = 'TOTAL'
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[1].text = str(deal['totals']['totalUnits'])
    total_row[1].paragraphs[0].runs[0].bold = True
    total_row[2].text = ''
    total_row[3].text = f"{deal['totals']['grossValue']:,.0f}"
    total_row[3].paragraphs[0].runs[0].bold = True
    for cell in total_row:
        set_cell_shading(cell, 'E2EFDA')
    
    doc.add_paragraph()
    
    # 3. Deal Terms
    doc.add_heading('3. Deal Terms', level=1)
    
    terms = doc.add_paragraph()
    terms.add_run('• Bulk Discount: ').bold = True
    terms.add_run(f"10% ({format_aed(deal['totals']['discount10pct'])})\n")
    terms.add_run('• Referral Commission: ').bold = True
    terms.add_run(f"9% ({format_aed(deal['totals']['commission9pct'])})\n")
    terms.add_run('• Payment Structure: ').bold = True
    terms.add_run('100% Upfront Cash Payment\n')
    terms.add_run('• Title Transfer: ').bold = True
    terms.add_run('Immediate upon payment completion\n')
    terms.add_run('• Total Savings: ').bold = True
    terms.add_run(f"19% ({format_aed(deal['totals']['totalSavings'])})")
    
    # 4. Investment Calculator
    doc.add_heading('4. Investment Calculator', level=1)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    calc_data = [
        ('Gross Portfolio Value', format_aed(deal['totals']['grossValue'])),
        ('Less: Bulk Discount (10%)', f"({format_aed(deal['totals']['discount10pct'])})"),
        ('Less: Referral Commission (9%)', f"({format_aed(deal['totals']['commission9pct'])})"),
        ('Your Investment', format_aed(deal['totals']['investorPays'])),
        ('Day-1 Equity Created', f"{format_aed(deal['totals']['totalSavings'])} (19%)"),
    ]
    
    for i, (label, value) in enumerate(calc_data):
        row = table.rows[i].cells
        row[0].text = label
        row[1].text = value
        if i >= 3:
            row[0].paragraphs[0].runs[0].bold = True
            row[1].paragraphs[0].runs[0].bold = True
            set_cell_shading(row[0], 'E2EFDA')
            set_cell_shading(row[1], 'E2EFDA')
    
    doc.add_paragraph()
    
    # 5. Rental Yield Projection
    doc.add_heading('5. Projected Returns (Conservative 7% Yield)', level=1)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    yield_data = [
        ('Metric', 'Amount'),
        ('Annual Gross Income', format_aed(deal['yields']['conservative7pct']['annual'])),
        ('Monthly Cash Flow', format_aed(deal['yields']['conservative7pct']['monthly'])),
    ]
    
    for i, (label, value) in enumerate(yield_data):
        row = table.rows[i].cells
        row[0].text = label
        row[1].text = value
        if i == 0:
            row[0].paragraphs[0].runs[0].bold = True
            row[1].paragraphs[0].runs[0].bold = True
            set_cell_shading(row[0], 'D9E2F3')
            set_cell_shading(row[1], 'D9E2F3')
    
    doc.add_paragraph()
    
    # 6. Legal Framework
    doc.add_heading('6. Legal and Implementation Framework', level=1)
    
    legal = doc.add_paragraph()
    legal.add_run('• All investments to be governed under UAE law.\n')
    legal.add_run('• Definitive agreements to include Term Sheet, MOA and investment agreement.\n')
    legal.add_run('• One Development will act as developer and manager under approved RERA and DLD guidelines.\n')
    legal.add_run('• Immediate title transfer upon payment completion.')
    
    # 7. Disclaimer
    doc.add_heading('7. Disclaimer', level=1)
    
    disclaimer = doc.add_paragraph()
    disclaimer.add_run('All information, figures, and projections in this document are ')
    disclaimer.add_run('indicative only').bold = True
    disclaimer.add_run(' and subject to confirmation following completion of due diligence, regulatory approvals, and final investor agreement.\n\n')
    disclaimer.add_run('Nothing herein constitutes an offer or commitment until formal binding documentation is executed by the Parties.')
    
    # About One Development
    doc.add_heading('8. About One Development', level=1)
    
    about = doc.add_paragraph()
    about.add_run('ONE Development is a real estate developer dedicated to creating lifestyle-driven, future-ready destinations in high-potential markets. Each project is approached as a holistic living experience, harmonizing modern comfort standards with thoughtful planning.\n\n')
    about.add_run('Portfolio Locations: ').bold = True
    about.add_run('Dubai • Abu Dhabi • Ras Al Khaimah • New Cairo • Riyadh • Athens')
    
    doc.add_paragraph()
    
    # Signature
    sig = doc.add_paragraph()
    sig.add_run('Prepared by:\n').bold = True
    sig.add_run('One International Real Estate Development LLC\n').bold = True
    sig.add_run('Dubai – United Arab Emirates\n')
    sig.add_run(f'Date: {datetime.now().strftime("%d %B %Y")}')
    
    # QR Code note
    doc.add_paragraph()
    qr_note = doc.add_paragraph()
    qr_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_note.add_run('Scan QR Code for Interactive Presentation\n').bold = True
    qr_note.add_run('https://investor-pitch.onestrategy.app')
    
    doc.save('External-Investor-Proposal-Laguna-Residence.docx')
    print('✅ Created: External-Investor-Proposal-Laguna-Residence.docx')


def create_internal_memo():
    """Create Internal Approval Memo in One Development format"""
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Header info
    header = doc.add_paragraph()
    header.add_run('Ref. No.: ').bold = True
    header.add_run('ONE/STR/2026/INV-001\n')
    header.add_run('Date: ').bold = True
    header.add_run(f'{datetime.now().strftime("%d %B %Y")}\n\n')
    header.add_run('To: ').bold = True
    header.add_run('Senior Management / Chairman\n')
    header.add_run('From: ').bold = True
    header.add_run('Strategy Department\n')
    header.add_run('Subject: ').bold = True
    run = header.add_run('Laguna Residence Bulk Sale – AED 893M Portfolio Approval Request')
    run.bold = True
    run.underline = True
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    summary = doc.add_paragraph()
    summary.add_run(f'This memorandum requests approval for a bulk sale transaction of {deal["totals"]["totalUnits"]} units at Laguna Residence totaling {format_aed(deal["totals"]["grossValue"])} in gross value. ')
    summary.add_run(f'The proposed terms include a 9% referral commission and 10% bulk discount, resulting in net proceeds of {format_aed(deal["totals"]["investorPays"])}.')
    
    # Deal Structure
    doc.add_heading('2. Deal Structure', level=1)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    structure_data = [
        ('Parameter', 'Value'),
        ('Gross Portfolio Value', format_aed(deal['totals']['grossValue'])),
        ('Total Units', str(deal['totals']['totalUnits'])),
        ('Average Unit Price', format_aed(deal['totals']['grossValue'] // deal['totals']['totalUnits'])),
        ('Net Proceeds', format_aed(deal['totals']['investorPays'])),
    ]
    
    for i, (label, value) in enumerate(structure_data):
        row = table.rows[i].cells
        row[0].text = label
        row[1].text = value
        if i == 0:
            row[0].paragraphs[0].runs[0].bold = True
            row[1].paragraphs[0].runs[0].bold = True
            set_cell_shading(row[0], 'D9E2F3')
            set_cell_shading(row[1], 'D9E2F3')
    
    doc.add_paragraph()
    
    # Proposed Terms
    doc.add_heading('3. Proposed Terms', level=1)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    terms_data = [
        ('Term', 'Rate', 'Impact'),
        ('Referral Commission', '9%', f'({format_aed(deal["totals"]["commission9pct"])})'),
        ('Bulk Discount', '10%', f'({format_aed(deal["totals"]["discount10pct"])})'),
        ('Payment Structure', '100% Upfront', 'Immediate liquidity'),
        ('Total Concession', '19%', f'({format_aed(deal["totals"]["totalSavings"])})'),
    ]
    
    for i, row_data in enumerate(terms_data):
        row = table.rows[i].cells
        for j, value in enumerate(row_data):
            row[j].text = value
            if i == 0:
                row[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row[j], 'D9E2F3')
            if i == 4:
                row[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row[j], 'E2EFDA')
    
    doc.add_paragraph()
    
    # Unit Allocation
    doc.add_heading('4. Unit Allocation', level=1)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Unit Type', 'Count', 'Avg. Price', 'Subtotal']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    unit_data = [
        ('Studio', deal['unitMix']['studio']['count'], deal['unitMix']['studio']['avgPrice'], deal['unitMix']['studio']['subtotal']),
        ('1BR', deal['unitMix']['oneBedroom']['count'], deal['unitMix']['oneBedroom']['avgPrice'], deal['unitMix']['oneBedroom']['subtotal']),
        ('2BR', deal['unitMix']['twoBedroom']['count'], deal['unitMix']['twoBedroom']['avgPrice'], deal['unitMix']['twoBedroom']['subtotal']),
    ]
    
    for i, (unit_type, count, avg, total) in enumerate(unit_data, 1):
        row = table.rows[i].cells
        row[0].text = unit_type
        row[1].text = str(count)
        row[2].text = format_aed(avg)
        row[3].text = format_aed(total)
    
    total_row = table.rows[4].cells
    total_row[0].text = 'Total'
    total_row[1].text = str(deal['totals']['totalUnits'])
    total_row[2].text = ''
    total_row[3].text = format_aed(deal['totals']['grossValue'])
    for cell in total_row:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E2EFDA')
    
    note = doc.add_paragraph()
    note.add_run('Data sourced from Salesforce Unit__c (Base_Price__c field) — January 30, 2026').italic = True
    note.paragraph_format.space_before = Pt(6)
    
    # Financial Analysis
    doc.add_heading('5. Financial Analysis', level=1)
    
    calc = doc.add_paragraph()
    calc.add_run('Proceeds Breakdown:\n').bold = True
    calc.add_run(f'Gross Value:                    {format_aed(deal["totals"]["grossValue"])}\n')
    calc.add_run(f'Less: Referral Commission (9%)  ({format_aed(deal["totals"]["commission9pct"])})\n')
    calc.add_run(f'Less: Bulk Discount (10%)       ({format_aed(deal["totals"]["discount10pct"])})\n')
    calc.add_run('─' * 45 + '\n')
    run = calc.add_run(f'NET PROCEEDS:                   {format_aed(deal["totals"]["investorPays"])}')
    run.bold = True
    
    # Strategic Rationale
    doc.add_heading('6. Strategic Rationale', level=1)
    
    benefits = doc.add_paragraph()
    benefits.add_run('Benefits:\n').bold = True
    benefits.add_run(f'1. Immediate Liquidity — {format_aed(deal["totals"]["investorPays"])} cash injection\n')
    benefits.add_run('2. Risk Mitigation — Eliminates 24-36 month sales cycle\n')
    benefits.add_run('3. Operational Efficiency — Single transaction vs. 632 individual sales\n')
    benefits.add_run('4. Market Positioning — Demonstrates institutional investor interest\n')
    benefits.add_run(f'5. Remaining Inventory — 307 units (~AED 901M) remain for retail sales')
    
    # Recommendation
    doc.add_heading('7. Recommendation', level=1)
    
    rec = doc.add_paragraph()
    run = rec.add_run('APPROVE ')
    run.bold = True
    rec.add_run('this transaction based on:\n')
    rec.add_run('• Strong NPV compared to alternative scenarios\n')
    rec.add_run('• Immediate liquidity benefits\n')
    rec.add_run('• Risk mitigation through 100% upfront payment\n')
    rec.add_run('• Operational efficiency gains\n')
    rec.add_run('• Strategic portfolio optimization')
    
    doc.add_paragraph()
    
    # Approval Signatures
    doc.add_heading('8. Approval Signatures', level=1)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    sig_headers = ['Role', 'Name', 'Signature', 'Date']
    for i, header in enumerate(sig_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    roles = ['Strategy Director', 'CFO', 'CEO', 'Chairman']
    for i, role in enumerate(roles, 1):
        table.rows[i].cells[0].text = role
    
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.add_run('Prepared by: ').bold = True
    footer.add_run('Strategy Department\n')
    footer.add_run('Classification: ').bold = True
    footer.add_run('CONFIDENTIAL - INTERNAL USE ONLY')
    
    doc.save('Internal-Approval-Memo-Laguna-Bulk-Sale.docx')
    print('✅ Created: Internal-Approval-Memo-Laguna-Bulk-Sale.docx')


if __name__ == '__main__':
    create_external_proposal()
    create_internal_memo()
    print('\n📄 Word documents created successfully!')
